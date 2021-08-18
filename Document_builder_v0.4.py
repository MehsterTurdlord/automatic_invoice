from docx import Document
from pprint import pprint
# For version 4, maintain the original formatting and add that code to make the amount word.

def get_invoice_number():
    '''
    Reads the configuration document and extracts the invoice number, then iterating and returning it.
    Returns: [int] invoice number + 1
    '''
    Invoice_No = 0
    I_doc = Document('Invoice.docx')
    Invoice_No = int(I_doc.paragraphs[0].text)

    I_doc.save('Invoice.docx')
    
    print("Success ! Got Invoice #")
    return Invoice_No

def set_invoice_number(Invoice_No):
    '''
    After a successful run, it takes the invoice number and writes the configuration file and replaces the old with the new.
    Returns: nothing, but perhaps a "success" note.
    '''
    I_doc = Document('Invoice.docx')
    I_doc.paragraphs[0].text = str(Invoice_No + 1)

    I_doc.save('Invoice.docx')
    print("Success ! Set Invoice #")
    

def main():
    
    doc = Document('original_file.docx')
    C_asked = False
    n = int(input('How many of the same INVOICES for the SAME customer ? (type a number) ') or 1)
    
    # Collecting
    if C_asked is False:
        C_Name = str(input('What is the CUSTOMER\'S COMPANY NAME '))
        C_Telephone = str(input('What is the CUSTOMER\'S TELEPHONE # ? '))
        C_Person = str(input('Who is the CONTACT PERSON ? '))
        C_TRN =  str(input('What is the CUSTOMER\'S TRN ? '))

        print("Success ! Got customer details")
  
        C_asked = True
        
    while n >= 1:
        D_Invoice = get_invoice_number()

        Description = str(input('What is the DESCRIPTION ? '))
        Quantity = float(input('What is the QUANTITY ? '))
        Unit_Amount = float(input('What is the UNIT AMOUNT ? '))
        VAT_Amount = Unit_Amount*Quantity*0.05
        Total_Amount = Unit_Amount*Quantity*1.05

        print("Success ! Got contract details")


        # Editing
        doc.paragraphs[5].text = ''.join(['Company name: ', C_Name])

        doc.tables[0].cell(0,0).text = ''.join(['Telephone No: ', C_Telephone])
        doc.tables[0].cell(0,1).text = ''.join(['Contact Person: ' , C_Person])
        doc.tables[0].cell(0,2).text = ''.join(['Customer TRN: ' , C_TRN])
        doc.tables[0].cell(1,0).text = ''.join(['Invoice No: ' , str(D_Invoice)])

        doc.tables[1].cell(1,0).text = '01'
        doc.tables[1].cell(1,1).text = '01'
        doc.tables[1].cell(1,2).text = Description
        doc.tables[1].cell(1,3).text = str(int(Quantity))
        doc.tables[1].cell(1,4).text = str(int(Unit_Amount))
        doc.tables[1].cell(1,5).text = str(round(Unit_Amount - int(Unit_Amount),2)).split(".")[-1]
        doc.tables[1].cell(1,6).text = str(int(VAT_Amount))
        doc.tables[1].cell(1,7).text = str(round(VAT_Amount - int(VAT_Amount),2)).split(".")[-1]
        doc.tables[1].cell(1,8).text = str(int(Total_Amount))
        doc.tables[1].cell(1,9).text = str(round(Total_Amount - int(Total_Amount),2)).split(".")[-1]

        doc.tables[1].cell(9,8).text = ''.join(['AED: ', str(int(Quantity*Unit_Amount))])
        doc.tables[1].cell(9,9).text = str(round(Quantity*Unit_Amount - int(Quantity*Unit_Amount),2)).split(".")[-1]
        doc.tables[1].cell(10,8).text = ''.join(['AED: ', str(int(VAT_Amount))])
        doc.tables[1].cell(10,9).text = str(round(VAT_Amount - int(VAT_Amount),2)).split(".")[-1]
        doc.tables[1].cell(11,8).text = ''.join(['AED: ', str(int(Total_Amount))])
        doc.tables[1].cell(11,9).text = str(round(Total_Amount - int(Total_Amount),2)).split(".")[-1]

        # Saving
        set_invoice_number(D_Invoice)
        doc.save(' '.join([str(D_Invoice),  C_Name])  + '.docx')
        n = n - 1
        print("Success ! Created invoice !")
        
    C_asked = False
    
if __name__ == '__main__':
    print('Starting')
    main()
else:
    print('hai')

