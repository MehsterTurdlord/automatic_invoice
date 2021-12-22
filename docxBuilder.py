from docx import Document


def build_invoice(values):
    """
    Either this or the main should be the one that builds the invoice.

    Parameters
    -------
    values: Dict
        'cNam': str, 'cTel': int, 'cPer': str,  
        'cTRN': int, 'dInv': int, 'products': list,
        'totalValues': float, 'totalWords: str

    products: list
        'pDes': str, 'pQua': float, 'pCos': float, 
        'base': float, 'VAT': float, 'total': float

    Returns
    -------
    None
    """

    doc = Document('original_file.docx')

    
    doc.paragraphs[7].text       = 'Company name: '   + values['cNam']
    doc.tables[0].cell(0,0).text = 'TELEPHONE NO: '   + values['cTel']
    doc.tables[0].cell(1,0).text = 'CONTACT PERSON: ' + values['cPer']
    doc.tables[0].cell(2,0).text = 'CUSTOMER TRN: '   + values['cTRN']
    doc.tables[0].cell(0,1).text = 'INVOICE No: '     + values['dInv']
    doc.tables[0].cell(1,1).text = 'DATE: '           + values['date']

    #TODO: Fix the pQua in this for loop properly.
    for n in range(1, len( values['products']) + 1 ):
        doc.tables[1].cell(n,0).text = str(n)
        doc.tables[1].cell(n,2).text = values['products'][n-1]['pDes'];
        doc.tables[1].cell(n,3).text = str(int(values['products'][n-1]['pQua']));
        doc.tables[1].cell(n,4).text = values['products'][n-1]['base'][0];
        doc.tables[1].cell(n,5).text = values['products'][n-1]['base'][1];
        doc.tables[1].cell(n,6).text = values['products'][n-1]['VAT'][0];
        doc.tables[1].cell(n,7).text = values['products'][n-1]['VAT'][1];
        doc.tables[1].cell(n,8).text = values['products'][n-1]['total'][0];
        doc.tables[1].cell(n,9).text = values['products'][n-1]['total'][1];

    # The total base price
    doc.tables[1].cell(9,8).text  = 'AED: ' + values['totalValues'][0]['base'][0]
    doc.tables[1].cell(9,9).text  = values['totalValues'][0]['base'][1]
    
    # The total VAT price
    doc.tables[1].cell(10,8).text = 'AED: ' + values['totalValues'][0]['VAT'][0]
    doc.tables[1].cell(10,9).text = values['totalValues'][0]['VAT'][1]

    # The total total price    
    doc.tables[1].cell(11,8).text = 'AED: ' + values['totalValues'][0]['total'][0]
    doc.tables[1].cell(11,9).text = values['totalValues'][0]['total'][1]

    # The total total in words
    doc.tables[1].cell(12,0).text = values['totalWords'] 

    # Saving
    doc.save(values['dInv'] + ' ' + values['cNam'] + '.docx')

    print("Success ! Created invoice !")
    
    return None
