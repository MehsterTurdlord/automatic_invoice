from docx import Document
from datetime import datetime
from pprint import pprint
import configparser
import os

def create_config_file():
    """Creates a configuration file."""

    write_config = configparser.ConfigParser();
    write_config.add_section("Section_1");
    write_config.set("Section_1", "invoice", "9400");

    cfgfile = open("config.ini", 'w');
    write_config.write(cfgfile);
    cfgfile.close();

    return None


def get_invoice_number():
    '''
    Reads the configuration file and extracts the invoice number, then returns after iteration.

    Parameters
    -------
    invoice_no: str
        The invoice number last used

    Returns
    -------
    invoice_no: str
        an interated invoice_no
    '''
    invoice_no = "";

    try:
        cfg = configparser.ConfigParser();
        cfg.read("config.ini");
        invoice_no = cfg.get("Section_1", "invoice");
        invoice_no = str( (int(invoice_no) + 1 ));

        print(invoice_no);
        print("Success ! Got Invoice #.");
  
    except:
        print("Failure ! Did not get Invoice #.");

    return invoice_no


def set_invoice_number(invoice_no):
    '''
    After a successful run, it takes the invoice number and writes the configuration file and replaces the old with the new.
    Returns: nothing, but perhaps a "success" note.

    Parameters
    -------
    I_doc: Document
        The open document that saves the invoice number.

    Invoice_No: str
        This is converted to int, iterated, and then saved into the file.

    Returns
    ------
    None
    '''

    try:
        cfg = configparser.ConfigParser();
        cfg.set("Section_1", "invoice", invoice_no);

        cfg = open("config.ini", 'w');
        cfg.write(cfg);
        cfg.close();

        print("Success ! Set Invoice #.")

    except:
        print("Failure ! Did not set Invoice #.");

    return None


def build_invoice(values):
    """
    Either this or the main should be the one that builds the invoice.

    Parameters
    -------
    values: Dict
        'cNam': str, 'cTel': int, 'cPer': str,  
        'cTRN': int, 'dInv': int, 'products': list,
        'totalValues': float, 'totalWords: str

    products: list
        'pDes': str, 'pQua': float, 'pCos': float, 
        'base': float, 'VAT': float, 'total': float

    Returns
    -------
    None
    """

    doc = Document('original_file.docx')

    
    doc.paragraphs[7].text       = 'Company name: '   + values['cNam']
    doc.tables[0].cell(0,0).text = 'TELEPHONE NO: '   + values['cTel']
    doc.tables[0].cell(1,0).text = 'CONTACT PERSON: ' + values['cPer']
    doc.tables[0].cell(2,0).text = 'CUSTOMER TRN: '   + values['cTRN']
    doc.tables[0].cell(0,1).text = 'INVOICE No: '     + values['dInv']
    doc.tables[0].cell(1,1).text = 'DATE: '           + values['date']

    #TODO: Fix the pQua in this for loop properly.
    for n in range(1, len( values['products']) + 1 ):
        doc.tables[1].cell(n,0).text = str(n)
        doc.tables[1].cell(n,2).text = values['products'][n-1]['pDes'];
        doc.tables[1].cell(n,3).text = str(int(values['products'][n-1]['pQua']));
        doc.tables[1].cell(n,4).text = values['products'][n-1]['base'][0];
        doc.tables[1].cell(n,5).text = values['products'][n-1]['base'][1];
        doc.tables[1].cell(n,6).text = values['products'][n-1]['VAT'][0];
        doc.tables[1].cell(n,7).text = values['products'][n-1]['VAT'][1];
        doc.tables[1].cell(n,8).text = values['products'][n-1]['total'][0];
        doc.tables[1].cell(n,9).text = values['products'][n-1]['total'][1];

    # The total base price
    doc.tables[1].cell(9,8).text  = 'AED: ' + values['totalValues'][0]['base'][0]
    doc.tables[1].cell(9,9).text  = values['totalValues'][0]['base'][1]
    
    # The total VAT price
    doc.tables[1].cell(10,8).text = 'AED: ' + values['totalValues'][0]['VAT'][0]
    doc.tables[1].cell(10,9).text = values['totalValues'][0]['VAT'][1]

    # The total total price    
    doc.tables[1].cell(11,8).text = 'AED: ' + values['totalValues'][0]['total'][0]
    doc.tables[1].cell(11,9).text = values['totalValues'][0]['total'][1]

    # The total total in words
    doc.tables[1].cell(12,0).text = values['totalWords'] 

    # Saving
    doc.save(values['dInv'] + ' ' + values['cNam'] + '.docx')

    print("Success ! Created invoice !")
    
    return None


def intConvert(num):
    """
    This transforms a constant from number format to one of letters and words.

    Parameters
    ------
    num: str/int
        This comes in as str and converted into an int.
    
    numEng: dict
        This is the unique set of numbers assigned an associated english word.

    Returns
    ------
    :str
        This is the constant in word form.
    """
    numEngA =  {
		0: 'zero', 1: 'one', 2: 'two', 3: 'three', 4: 'four',
		5: 'five', 6: 'six', 7: 'seven', 8: 'eight', 9: 'nine',
		10: 'ten', 11:'eleven', 12:'twelve', 13: 'thirteen',
		14: 'fourteen', 15: 'fifteen', 16: 'sixteen',
		17: 'seventeen', 18: 'eighteen', 19: 'nineteen'
		}
    
    numEngB = ['twenty', 'thirty', 'fourty', 'fifty',
       'sixty', 'seventy', 'eighty', 'ninety']
    
    num = int(num);

    if 1 <= num <= 19:
        return numEngA[num];

    elif 20 <= num <= 99:
        tens, units = divmod(num, 10);	    
        print('The tens: {}'.format(tens));
        print('The units: {}'.format(units));
        return numEngB[tens - 2] + '-' + numEngA[units] if units else numEngB[tens - 2];
    
    elif 100 <= num <= 999:
        hundreds, tens = divmod(num, 100);
        print('The hundreds: {}'.format(hundreds));
        print('The tens: {}'.format(tens));

        return intConvert(hundreds) + ' hundred' + ' and ' + intConvert(tens) if tens else intConvert(hundreds) + ' hundred';

    elif 1000 <= num <= 9999:
        thousands, hundreds = divmod(num, 1000);
        print('The thousands: {}'.format(thousands));
        print('The hundreds: {}'.format(hundreds));
        
        return intConvert(thousands) + ' thousand' + ' and ' + intConvert(hundreds) if hundreds else intConvert(thousands) + ' thousand';
	
    else:
	    print('Sorry, the number is outside of our boundaries !');


def decConvert(dec):
    """
	This is a number-word converter, but for decimals.
	
	Parameters
	-----
	dec:str
		This is the input value
	numEngA: dict
		A dictionary of values that are only up to single digits
	frstDP: int
            The first decimal place
	scndDP: int
            The second decimal place
	
	Returns
	-----
	:str
		This checks to see if there is a valid scndp, i.e., not zero,
		and then then returns a valid decmial value in English format.
	"""
    numEngA =  {
		0: 'zero', 1: 'one', 2: 'two', 3: 'three', 4: 'four',
		5: 'five', 6: 'six', 7: 'seven', 8: 'eight', 9: 'nine',
		}
    numEngB = {
        1: 'ten', 2: 'twenty', 3: 'thirty', 4: 'fourty',
		5: 'fifty', 6: 'sixty', 7: 'seventy', 8: 'eighty', 9: 'ninety',
        }

    frstDP = int(dec[0]);
    scndDP = int(dec[1]);
    return ' and ' + numEngA[frstDP] + ' ' + numEngA[scndDP] if not scndDP else ' and ' + numEngB[frstDP]


def minimum_format_ensurer(products):
    """
    This is a function to ensure that there is at least 2 digits to each input.

    It adds a leading or an ending zero to each single-digit entry,
    all values will be turned to str or tuples containing str to maintain the format.

    Any item not of int or float will be skipped.

    str.zfill() is a good alternative, but not so much for decimals.

    Parameters
    -------

    h: dict
        this is the dict of each product.

    k: str
        This is the key of each dict entry.

    i: str/int/float
        This is the value of each entry, which does not stay standard.

    Returns
    -------
    : list
    """
    for h in products:
        for k, i in h.items():
            if k == 'pQua':
                continue;
                
            if type(i) is int and len(str(i)) == 1:
                #print('is int !')
                i = '0' + str(i)

            elif type(i) is float:
                #print('is float !')
                AED = str(i).split('.')[0]
                FIL = str(i).split('.')[-1]

                if len(AED) == 1:
                    AED = '0' + AED

                if len(FIL) == 1:
                    FIL = FIL + '0'

                i = (AED, FIL)
                #print(AED, FIL)

            else:
                #print('is undesirable !')
                continue

            h[k] = i;

    return products


def totaler(products):
    """Totals the total value of each product."""
    totalDict = {'base': 0, 'VAT': 0, 'total': 0};

    for h in products:
        totalDict['base'] += h['base'];
        totalDict['VAT'] += h['VAT'];
        totalDict['total'] += h['total'];

    return totalDict;


def splitter(h):
    """ Splits dictionary numbers by the decimal point."""
    if type(h) is dict:
        for k, i in h.items():
            h[k] = str(i).split('.');

    if type(h) is list:
        for n in range(0, len(h)):
            h[n] = splitter(h[n])

    return h


def get_date():
    """Get's todays date from the datetime module in the requested format."""
    return datetime.today().date().strftime('%B %d %Y');


def get_values():
    """Gets the values from the CLI."""
    

def main(GUIDetails = None):
    """
    First, input and get_invoice_number().
    Then we prepare the values with calculations and minimum_num_format().
    Finally, we build with build_invoice() and end with set_invoice_number().

    Parameters
    -------
    cNam: str
    cTel: str
    cPer: str
    cTRN: str
    dInv int
        This is received from get_invoice_number() and it is used to update through set_invoice_number()
    
    products: list
        this is a list made of dicts, the following items - range ending at 'values', exclusively - make up each
         dict in order to allow for multiple items.

    pDes: str
    pQua: float
    pCos: float
    base: float
        The pQua * the pCos
    VAT: float
        The Base_Amount * 0.05
    total: float
        The Base_Amount * 1.05

    values: dict
        {'cNam': cNam, 'cTel': cTel, 'cPer': cPer,  
         'cTRN': cTRN, 'dInv': dInv, 'products': products, }

    k: dict
        this is the temporary pointer to the dicts made during the loops.

    Returns
    -------
    None
    """
    
    if (os.path.isfile('config.ini')) is False:
        create_config_file();

    print('Getting invoice number...')

    dInv = get_invoice_number()
    products = []

    # Depreciate this.
    try:
        cNam = str( GUIDetails[0] );
        cTel = str( GUIDetails[1] );
        cPer = str( GUIDetails[2] );
        cTRN = str( GUIDetails[3] );
        
        for n in range(4, len(GUIDetails) + 1 ):
            h = {}

            h['pDes'] = str(   GUIDetails[n][0] );
            h['pQua'] = float( GUIDetails[n][1] );
            h['pCos'] = float( GUIDetails[n][2] );

            products.append(h);
    
    except:   
        cNam = str( input( 'What is the CUSTOMER\'S COMPANY NAME ? ' ) or 'University of Wollongong in Dubai')
        cTel = str( input( 'What is the CUSTOMER\'S TELEPHONE # ? '  ) or '971-56-1322345'                   )
        cPer = str( input( 'Who is the CONTACT PERSON ? '            ) or 'Mr Wollongong'                    )
        cTRN = str( input( 'What is the CUSTOMER\'S TRN ? '          ) or '12381239018'                      )

        print("Success ! Received customer details.")

        while True:
            h = {}
            h['pDes'] = str(   input( 'What is the Product Description? ' ) or 'Printers');
            h['pQua'] = float( input( 'How much is the Quantity ? '       ) or 5         );
            h['pCos'] = float( input( 'How much is the Unit Amount ? '    ) or 50        );

            products.append(h);
            
            if   str( input( 'No more products ? (Y/N)' ) or 'N' ).upper().split()[0] == 'Y':
                break;
            else:
                continue;

        print("Success ! Received product details.")

    
    print('Calculating values...')
    for h in products:
        h['base']  = round( h['pQua'] * h['pCos'], 2 );
        h['VAT']   = round( h['base'] * 0.05     , 2 );
        h['total'] = round( h['base'] + h['VAT'] , 2 );

    print('Ensuring minimum format...')
    values = {'cNam': cNam, 'cTel': cTel, 'cPer': cPer,  
         'cTRN': cTRN, 'dInv': dInv, 'products': products, }

    pprint(values);
    
    print('Converting numbers into words...');
    
    values['totalValues'] = totaler(  values['products'    ] );
    values['totalValues'] = minimum_format_ensurer( [ values['totalValues'] ] );
    #values['totalValues'] = splitter( values['totalValues' ] );

    # splitting these too
    values['products'] = minimum_format_ensurer(values['products']);
    #values['products'] = splitter( values['products'] )

    values['totalWords'] = ''.join(['AED: ', 
        intConvert( values['totalValues'][0]['total'][0] ), 
        decConvert( values['totalValues'][0]['total'][1] ),
        ' FILS', ' only']).upper();
    #values['totalWords'] = 'WIP'
    
    print('Acquiring dates...')
    values['date'] = get_date()

    print('Building invoices...')
    build_invoice(values)

    print('Setting invoice number...');
    set_invoice_number(dInv);


    return None


if __name__ == '__main__':
    print('Starting')
    main()
    
else:
    print('hai')
